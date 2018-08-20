#!/usr/bin/env python
# -*- coding: utf-8 -*-

# 根据已经存在的题库Word文档，自动导入到Excel表
import xlwt
import os
import tkinter.filedialog
import docx
from docx import Document
import tkinter.messagebox
from tkinter import *

# <editor-fold desc="***获得选中的文件路劲***">
default_dir = r"默认路径XBW"  # 默认路径。
# fname 为选中的文件路径。
fname = tkinter.filedialog.askopenfilename(title=u'选择文件XBW', initialdir=(os.path.expanduser((default_dir))))
# </editor-fold>

# 获取文件内容
document = Document(fname)
fstr = ""
for a in document.paragraphs:
    fstr = fstr + a.text + '\n'

# 匹配所有题目
Tpattern = re.compile(r'(\d+\..*?)A\.', re.S)
Tresult = re.findall(Tpattern, fstr)

# 将匹配到的所有题目写入Excel表
book = xlwt.Workbook(encoding='utf-8')
sheel = book.add_sheet("题库一")
i = 1
for a in Tresult:
    sheel.write(i, 5, a)
    i+=1

# 匹配选项
Xpattern = re.compile(r'[A-F]\..*')
Xresult = re.findall(Xpattern, fstr)
# 将匹配的选项写入到Excel表中
ans =""
j = 1
for a in Xresult:
    if (a[0]=="A" and len(ans)!=0):
        rans = ans[:-2]
        sheel.write(j, 6, rans)
        j+=1;
        ans=""
    ans = ans + a + "|\n"

# 匹配正确答案？？？
for a in document.paragraphs:
    b = a.text
    if (len(b)!=0 and b[0]>='A' and b[0]<='F'):
        co = a.style.font.color.rgb
        tt = b + "   颜色：" + str(a.style.font.color.rgb)
        print(tt)
# 保存
path = r'E:\2363521277\FileRecv\AAA.xls'
ex = os.path.exists(path)
if ex:
    os.remove(path)
book.save(path)






